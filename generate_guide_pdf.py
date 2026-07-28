"""Generation du Guide Fonctionnel de la Plateforme KYC.

Deux formats de sortie a partir d'un contenu unique (GUIDE_CONTENT) :
    python generate_guide_pdf.py            # PDF (comportement historique)
    python generate_guide_pdf.py --docx     # Word (.docx)
    python generate_guide_pdf.py --both     # les deux
"""

import argparse
import os
from datetime import datetime

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import (
    BaseDocTemplate, Frame, PageTemplate, Paragraph, Spacer, Table,
    TableStyle, PageBreak, Image, NextPageTemplate
)
from reportlab.platypus.tableofcontents import TableOfContents

                     
BOA_GREEN       = colors.HexColor("#00965E")
BOA_GREEN_LIGHT = colors.HexColor("#E8F5E9")
BOA_DARK        = colors.HexColor("#0f172a")
BOA_BLUE        = colors.HexColor("#1B2A4A")
BOA_SLATE       = colors.HexColor("#64748b")
BOA_WHITE       = colors.white
BOA_BORDER      = colors.HexColor("#e2e8f0")
BOA_GRAY        = colors.HexColor("#f1f5f9")

                                                                                    
HEX_GREEN  = "00965E"
HEX_BLUE   = "1B2A4A"
HEX_DARK   = "0F172A"
HEX_SLATE  = "64748B"
HEX_GRAY   = "F1F5F9"
HEX_WHITE  = "FFFFFF"

LOGO_PATH = os.path.join(os.getcwd(), "media", "images", "boa_logo.png")
SCREENSHOTS_DIR = os.path.join(os.getcwd(), "media", "screenshots")

DOC_TITLE = "Guide d'Utilisation"
DOC_SUBTITLE = "Plateforme Notation KYC"
DOC_BASELINE = "Manuel Fonctionnel illustre"
DOC_VERSION = "2.0"
FOOTER_TEXT = "BOA Group - Confidentiel | Guide d'utilisation KYC"
HEADER_TEXT = "GUIDE D'UTILISATION - PLATEFORME KYC"

                                                  
os.makedirs(SCREENSHOTS_DIR, exist_ok=True)


                                                                               
                                                                        
                                                                    
                                               
                                                                               
GUIDE_CONTENT = [
    ("h1", "1. Contexte et Objectifs"),
    ("h2", "Contexte et problématique"),
    ("p", "Les filiales du Groupe Bank of Africa font face à des insuffisances persistantes en matière "
          "de complétude et de fiabilisation des données KYC. Malgré les efforts des équipes "
          "opérationnelles, la mise à jour des dossiers reste difficile, notamment lorsque les clients "
          "sont difficilement joignables."),
    ("p", "Or, les documents scannés reçus, archivés dans des bases de données issues de projets comme "
          "INGEC ou DOCKFLOW, contiennent souvent les informations manquantes : numéros de pièces "
          "d'identité, dates de naissance, nationalités, adresses."),
    ("h2", "La solution : la Plateforme KYC et le module Screening KYC ID"),
    ("p", "La plateforme centralise le pilotage de la qualité des données KYC du Groupe. Le module "
          "Screening KYC ID extrait automatiquement les informations contenues dans les documents "
          "scannés, puis les rapproche des clients présentant des insuffisances de complétude KYC. "
          "Le module permet de :"),
    ("bullet", "Faciliter le travail de fiabilisation des dossiers clients."),
    ("bullet", "Réduire les risques de conformité."),
    ("bullet", "Améliorer la qualité globale du référentiel client."),
    ("p", "Le tout avec validation humaine avant toute mise à jour dans le corebanking."),
    ("h2", "Objectifs"),
    ("bullet", "Prioriser les dossiers pouvant être régularisés rapidement, en complément des campagnes "
               "de contact client."),
    ("bullet", "Identifier les informations disponibles dans les documents scannés."),
    ("bullet", "Limiter les erreurs de ressaisie grâce aux suggestions issues du document source."),
    ("bullet", "Réduire le volume de clients présentant des insuffisances de complétude KYC."),
    ("bullet", "Donner aux équipes une vision exploitable des correspondances entre documents scannés "
               "et dossiers clients."),
    ("bullet", "Contribuer à protéger le Groupe contre les risques de sanctions liés aux manquements KYC."),
    ("pagebreak",),

    ("h1", "2. Le Tableau de Bord"),
    ("p", "Dès votre connexion, vous arrivez sur le Tableau de Bord (Dashboard). "
          "Il vous donne une vision synthétique de la situation KYC de votre périmètre :"),
    ("shot", "dashboard.png", "Vue d'ensemble du Tableau de bord"),

    ("h1", "3. Gestion des Notations KYC"),
    ("h2", "Nouvelle Notation"),
    ("p", "Cet écran vous permet de créer ou mettre à jour le profil KYC d'un client."),
    ("shot", "nouvelle_notation.png", "Écran de saisie d'une Notation"),
    ("h2", "Historique des Notations"),
    ("p", "Permet de consulter toutes les notations passées, avec possibilité de recherche et de filtres."),
    ("shot", "historique.png", "Liste de l'historique des Notations"),
    ("pagebreak",),

    ("h1", "4. Suivis et Anomalies"),
    ("h2", "Clients en Anomalie"),
    ("p", "Liste détaillée des clients présentant un risque de conformité ou un dossier KYC incomplet. "
          "Vous pouvez exporter ces données."),
    ("shot", "clients_anomalie.png", "Écran des clients en anomalie"),
    ("h2", "Champs non renseignés"),
    ("p", "Identifiez rapidement les champs manquants dans les profils de vos clients pour organiser "
          "vos campagnes de mise à jour."),
    ("shot", "champs_non_renseignes.png", "Écran des champs manquants"),

    ("h1", "5. Conformité"),
    ("h2", "PPE et Comptes Spécifiques"),
    ("p", "Interfaces dédiées au suivi des Personnes Politiquement Exposées et des comptes à "
          "surveillance renforcée (ONG, Ambassades)."),
    ("shot", "conformite_ppe.png", "Suivi des PPE"),
    ("pagebreak",),

    ("h1", "6. Administration (PASS / DSI)"),
    ("h2", "Paramètres Système et Utilisateurs"),
    ("p", "Gestion complète des comptes utilisateurs, des habilitations et du paramétrage général."),
    ("shot", "parametres.png", "Gestion des Paramètres Système"),
    ("h2", "Règles de Qualité"),
    ("p", "Configuration dynamique des règles d'évaluation pour le scoring et les alertes."),
    ("shot", "regles_qualite.png", "Configuration des règles de qualité"),
    ("h2", "Audit"),
    ("p", "L'onglet Audit restitue la piste d'audit de la plateforme : connexions et déconnexions, "
          "échecs d'authentification, créations et modifications de comptes et d'habilitations, "
          "modifications des règles de qualité et traitements Screening KYC ID. Les événements sont "
          "filtrables par catégorie, acteur, filiale, période et statut, et exportables au format Excel. "
          "Le profil PASS consulte l'ensemble du Groupe ; un profil DSI ne consulte que les événements "
          "de sa propre filiale."),
    ("shot", "audit.png", "Piste d'audit de la plateforme"),
]


def get_styles():
    base = getSampleStyleSheet()
    styles = {
        "title": ParagraphStyle("title", parent=base["Normal"], fontName="Helvetica-Bold", fontSize=26, textColor=BOA_WHITE, leading=30, alignment=TA_LEFT),
        "heading1": ParagraphStyle("h1", parent=base["Normal"], fontName="Helvetica-Bold", fontSize=18, textColor=BOA_BLUE, leading=22, spaceBefore=20, spaceAfter=10),
        "heading2": ParagraphStyle("h2", parent=base["Normal"], fontName="Helvetica-Bold", fontSize=14, textColor=BOA_GREEN, leading=18, spaceBefore=15, spaceAfter=8),
        "body": ParagraphStyle("body", parent=base["Normal"], fontName="Helvetica", fontSize=10, textColor=BOA_DARK, leading=15, alignment=TA_JUSTIFY, spaceAfter=10),
        "bullet": ParagraphStyle("bullet", parent=base["Normal"], fontName="Helvetica", fontSize=10, textColor=BOA_DARK, leading=15, spaceAfter=5, bulletIndent=10, leftIndent=25),
        "footer": ParagraphStyle("footer", parent=base["Normal"], fontName="Helvetica", fontSize=8, textColor=BOA_SLATE, alignment=TA_CENTER),
        "toc_title": ParagraphStyle("toc_title", parent=base["Normal"], fontName="Helvetica-Bold", fontSize=20, textColor=BOA_BLUE, leading=24, spaceAfter=16),
        "toc1": ParagraphStyle("toc1", parent=base["Normal"], fontName="Helvetica-Bold", fontSize=11, textColor=BOA_BLUE, leading=18, spaceBefore=6),
        "toc2": ParagraphStyle("toc2", parent=base["Normal"], fontName="Helvetica", fontSize=10, textColor=BOA_DARK, leading=16, leftIndent=18),
    }
    return styles


def screenshot_path(filename):
    """Chemin de la capture si elle existe, sinon None."""
    path = os.path.join(SCREENSHOTS_DIR, filename)
    return path if os.path.exists(path) else None


def get_screenshot_flowable(filename, label="Capture d'écran"):
    """
    Returns an Image if the file exists, otherwise returns a grey Table as a placeholder.
    """
    path = screenshot_path(filename)
    if path:
        try:
                                                                                  
            return Image(path, width=16*cm, height=9*cm, kind='proportional')
        except Exception:
            pass

                                        
    data = [[f"[{label} - {filename}]\n\n(Placez l'image dans media/screenshots/{filename})"]]
    t = Table(data, colWidths=[16*cm], rowHeights=[7*cm])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), BOA_GRAY),
        ('TEXTCOLOR', (0,0), (-1,-1), BOA_SLATE),
        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('GRID', (0,0), (-1,-1), 1, BOA_BORDER),
        ('FONTNAME', (0,0), (-1,-1), 'Helvetica'),
        ('FONTSIZE', (0,0), (-1,-1), 10),
    ]))
    return t


                                                                               
            
                                                                               
class GuideDocTemplate(BaseDocTemplate):
    """BaseDocTemplate qui alimente le sommaire au fil de la construction."""

    def afterFlowable(self, flowable):
        if not isinstance(flowable, Paragraph):
            return
        style_name = flowable.style.name
        if style_name == "h1":
            self.notify("TOCEntry", (0, flowable.getPlainText(), self.page))
        elif style_name == "h2":
            self.notify("TOCEntry", (1, flowable.getPlainText(), self.page))


def generate_pdf(filename="Guide_Fonctionnel_Plateforme_KYC.pdf"):
    styles = get_styles()
    w, h = A4
    doc = GuideDocTemplate(filename, pagesize=A4, leftMargin=1.5*cm, rightMargin=1.5*cm, topMargin=2.5*cm, bottomMargin=2.0*cm)

    def draw_header_footer(canvas, doc):
        canvas.saveState()
                
        canvas.setFillColor(BOA_BLUE)
        canvas.rect(0, 0, w, 1*cm, fill=1, stroke=0)
        canvas.setStrokeColor(BOA_GREEN)
        canvas.setLineWidth(2)
        canvas.line(0, 1*cm, w, 1*cm)
        canvas.setFillColor(BOA_WHITE)
        canvas.setFont("Helvetica", 8)
        canvas.drawString(1*cm, 0.35*cm, FOOTER_TEXT)
        canvas.drawRightString(w - 1*cm, 0.35*cm, f"Page {doc.page}")

                               
        if doc.page > 1:
            canvas.setFillColor(BOA_GREEN)
            canvas.rect(0, h - 1.5*cm, w*0.6, 1.5*cm, fill=1, stroke=0)
            canvas.setFillColor(BOA_BLUE)
            canvas.rect(w*0.6, h - 1.5*cm, w*0.4, 1.5*cm, fill=1, stroke=0)

            canvas.setFillColor(BOA_WHITE)
            canvas.setFont("Helvetica-Bold", 10)
            canvas.drawString(1*cm, h - 0.9*cm, HEADER_TEXT)

            if os.path.exists(LOGO_PATH):
                try:
                    canvas.drawImage(LOGO_PATH, w - 4*cm, h - 1.3*cm, width=3*cm, height=1.1*cm, preserveAspectRatio=True, mask="auto")
                except Exception:
                    pass

        canvas.restoreState()

                
    def draw_cover(canvas, doc):
        canvas.saveState()
        canvas.setFillColor(BOA_GREEN)
        canvas.rect(0, h*0.5, w, h*0.5, fill=1, stroke=0)

                        
        canvas.setFillColor(colors.HexColor("#007A4A"))
        path = canvas.beginPath()
        path.moveTo(0, h*0.5)
        path.lineTo(w*0.7, h*0.5)
        path.lineTo(0, h*0.35)
        path.close()
        canvas.drawPath(path, fill=1, stroke=0)

        if os.path.exists(LOGO_PATH):
            try:
                canvas.drawImage(LOGO_PATH, w - 6*cm, h - 3*cm, width=5*cm, height=2*cm, preserveAspectRatio=True, mask="auto")
            except Exception:
                pass

        canvas.setFillColor(BOA_WHITE)
        canvas.setFont("Helvetica-Bold", 32)
        canvas.drawString(1.5*cm, h*0.7, DOC_TITLE)
        canvas.drawString(1.5*cm, h*0.63, DOC_SUBTITLE)

        canvas.setStrokeColor(BOA_WHITE)
        canvas.setLineWidth(3)
        canvas.line(1.5*cm, h*0.59, 10*cm, h*0.59)

        canvas.setFont("Helvetica", 14)
        canvas.drawString(1.5*cm, h*0.54, DOC_BASELINE)

        canvas.setFillColor(BOA_DARK)
        canvas.setFont("Helvetica-Bold", 12)
        canvas.drawString(1.5*cm, h*0.4, f"Date de mise à jour : {datetime.now().strftime('%d/%m/%Y')}")
        canvas.drawString(1.5*cm, h*0.35, f"Version : {DOC_VERSION}")

        canvas.restoreState()

    frame_cover = Frame(0, 0, w, h, id="cover")
    frame_normal = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id="normal")

    doc.addPageTemplates([
        PageTemplate(id="cover", frames=[frame_cover], onPage=draw_cover),
        PageTemplate(id="normal", frames=[frame_normal], onPage=draw_header_footer)
    ])

                                                                               
    story = [Spacer(1, 1), NextPageTemplate("normal"), PageBreak()]

                                                                      
    toc = TableOfContents()
    toc.levelStyles = [styles['toc1'], styles['toc2']]
    toc.dotsMinLevel = 0
    story.append(Paragraph("Sommaire", styles['toc_title']))
    story.append(toc)
    story.append(PageBreak())

    for item in GUIDE_CONTENT:
        kind = item[0]
        if kind == "h1":
            story.append(Paragraph(item[1], styles['heading1']))
        elif kind == "h2":
            story.append(Paragraph(item[1], styles['heading2']))
        elif kind == "p":
            story.append(Paragraph(item[1], styles['body']))
        elif kind == "bullet":
            story.append(Paragraph(item[1], styles['bullet'], bulletText="•"))
        elif kind == "shot":
            story.append(Spacer(1, 0.3*cm))
            story.append(get_screenshot_flowable(item[1], item[2]))
            story.append(Spacer(1, 0.5*cm))
        elif kind == "pagebreak":
            story.append(PageBreak())

                                                                             
    doc.multiBuild(story)
    print(f"PDF généré avec succès : {filename}")
    return filename


                                                                               
                     
                                                                               
def generate_docx(filename="Guide_Fonctionnel_Plateforme_KYC.docx"):
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    def hex_rgb(value):
        return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))

    def shade(element, hex_fill):
        """Applique une couleur de fond a une cellule ou un paragraphe."""
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_fill)
        element.append(shd)

    def add_field(paragraph, instruction, placeholder=None):
        """Insere un champ Word (ex. PAGE, TOC) dans un paragraphe.

        `placeholder` affiche un texte tant que le champ n'a pas ete calcule
        par Word (mise a jour par F9 ou a l'impression).
        """
        run = paragraph.add_run()
        begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
        instr.text = instruction
        end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
        run._r.append(begin); run._r.append(instr)
        if placeholder:
            sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
            run._r.append(sep)
            text = OxmlElement("w:t"); text.set(qn("xml:space"), "preserve")
            text.text = placeholder
            run._r.append(text)
        run._r.append(end)
        return run

    def styled_paragraph(text="", *, size=10, bold=False, color=HEX_DARK, align=None,
                         space_before=0, space_after=6, container=None, style=None):
        target = container if container is not None else doc
        p = target.add_paragraph(style=style) if style else target.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        if align is not None:
            p.alignment = align
        if text:
            run = p.add_run(text)
            run.font.name = "Calibri"
            run.font.size = Pt(size)
            run.bold = bold
            run.font.color.rgb = hex_rgb(color)
        return p

    def banner(text, hex_fill, *, size=26, color=HEX_WHITE, height_cm=None):
        """Bandeau de couleur pleine largeur (table 1x1 ombree)."""
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.rows[0].cells[0]
        cell.width = Cm(18)
        shade(cell._tc.get_or_add_tcPr(), hex_fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(14)
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(size)
        run.bold = True
        run.font.color.rgb = hex_rgb(color)
        if height_cm:
            table.rows[0].height = Cm(height_cm)
        return table

    doc = Document()

    section = doc.sections[0]
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2.8)
    section.bottom_margin = Cm(2)
    section.header_distance = Cm(1)

                      
    if os.path.exists(LOGO_PATH):
        cover_logo = styled_paragraph(align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=18)
        try:
            cover_logo.add_run().add_picture(LOGO_PATH, width=Cm(5))
        except Exception:
            pass

    banner(DOC_TITLE, HEX_GREEN, size=30)
    banner(DOC_SUBTITLE, HEX_BLUE, size=22)
    styled_paragraph(DOC_BASELINE, size=14, color=HEX_SLATE, space_before=18, space_after=30)
    styled_paragraph(f"Date de mise à jour : {datetime.now().strftime('%d/%m/%Y')}",
                     size=12, bold=True, color=HEX_DARK, space_after=4)
    styled_paragraph(f"Version : {DOC_VERSION}", size=12, bold=True, color=HEX_DARK, space_after=4)
    styled_paragraph(FOOTER_TEXT, size=9, color=HEX_SLATE, space_before=24)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

                                                            
                                                                   
    section.different_first_page_header_footer = True
    first_footer = section.first_page_footer.paragraphs[0]
    first_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = first_footer.add_run(FOOTER_TEXT)
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = hex_rgb(HEX_SLATE)

    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header_p.add_run(HEADER_TEXT)
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.bold = True
    run.font.color.rgb = hex_rgb(HEX_BLUE)
    if os.path.exists(LOGO_PATH):
        try:
            header_p.add_run("\t\t").font.size = Pt(9)
            header_p.add_run().add_picture(LOGO_PATH, width=Cm(3))
        except Exception:
            pass

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run(f"{FOOTER_TEXT}   |   Page ")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = hex_rgb(HEX_SLATE)
    page_run = add_field(footer_p, "PAGE")
    page_run.font.name = "Calibri"
    page_run.font.size = Pt(8)
    page_run.font.color.rgb = hex_rgb(HEX_SLATE)

                    
    styled_paragraph("Sommaire", size=20, bold=True, color=HEX_BLUE, space_after=14)
    toc_p = doc.add_paragraph()
    add_field(toc_p, r'TOC \o "1-2" \h \z \u',
              placeholder="Sommaire : clic droit sur cette zone puis « Mettre à jour les champs » (ou F9).")
    for run in toc_p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.color.rgb = hex_rgb(HEX_SLATE)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

                   
    for item in GUIDE_CONTENT:
        kind = item[0]

        if kind == "h1":
            p = styled_paragraph(item[1], size=18, bold=True, color=HEX_BLUE,
                                 space_before=18, space_after=8, style="Heading 1")
            borders = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "12")
            bottom.set(qn("w:space"), "4")
            bottom.set(qn("w:color"), HEX_GREEN)
            borders.append(bottom)
            p._p.get_or_add_pPr().append(borders)

        elif kind == "h2":
            styled_paragraph(item[1], size=14, bold=True, color=HEX_GREEN,
                             space_before=14, space_after=6, style="Heading 2")

        elif kind == "p":
            styled_paragraph(item[1], size=10, color=HEX_DARK,
                             align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=10)

        elif kind == "bullet":
            styled_paragraph(item[1], size=10, color=HEX_DARK,
                             space_after=4, style="List Bullet")

        elif kind == "shot":
            _, shot_file, label = item
            path = screenshot_path(shot_file)
            if path:
                p = styled_paragraph(align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
                try:
                    p.add_run().add_picture(path, width=Cm(16))
                except Exception:
                    path = None
            if not path:
                                                                
                table = doc.add_table(rows=1, cols=1)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                cell = table.rows[0].cells[0]
                cell.width = Cm(16)
                cell.height = Cm(7)
                shade(cell._tc.get_or_add_tcPr(), HEX_GRAY)
                cp = cell.paragraphs[0]
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cp.paragraph_format.space_before = Pt(40)
                cp.paragraph_format.space_after = Pt(40)
                run = cp.add_run(f"[{label} - {shot_file}]\n(Placez l'image dans media/screenshots/{shot_file})")
                run.font.name = "Calibri"
                run.font.size = Pt(10)
                run.font.color.rgb = hex_rgb(HEX_SLATE)
            styled_paragraph(label, size=9, color=HEX_SLATE,
                             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

        elif kind == "pagebreak":
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    doc.save(filename)
    print(f"Word généré avec succès : {filename}")
    return filename


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description="Genere le guide fonctionnel de la Plateforme KYC.")
    parser.add_argument("--docx", action="store_true", help="Genere uniquement la version Word (.docx)")
    parser.add_argument("--both", action="store_true", help="Genere les versions PDF et Word")
    parser.add_argument("-o", "--output", default="Guide_Fonctionnel_Plateforme_KYC",
                        help="Nom de fichier sans extension")
    args = parser.parse_args()

    if args.both:
        generate_pdf(f"{args.output}.pdf")
        generate_docx(f"{args.output}.docx")
    elif args.docx:
        generate_docx(f"{args.output}.docx")
    else:
        generate_pdf(f"{args.output}.pdf")
